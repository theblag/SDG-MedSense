import os
import uuid
import io
from typing import List, Dict, Any
from datetime import datetime
import pypdf
from docx import Document
import email
from email import policy
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def detect_document_type(self, text: str, filename: str) -> str:
        """Automatically detect document type based on content and filename"""
        text_lower = text.lower()
        filename_lower = filename.lower()
        
        # Keywords for different document types
        insurance_keywords = [
            'policy', 'coverage', 'premium', 'deductible', 'claim', 'insured', 'insurer',
            'policyholder', 'endorsement', 'exclusion', 'liability', 'property', 'casualty',
            'health insurance', 'life insurance', 'auto insurance', 'home insurance'
        ]
        
        legal_keywords = [
            'contract', 'agreement', 'terms and conditions', 'clause', 'section', 'article',
            'party', 'obligation', 'liability', 'breach', 'termination', 'jurisdiction',
            'legal', 'law', 'statute', 'regulation', 'compliance'
        ]
        
        financial_keywords = [
            'financial', 'revenue', 'profit', 'loss', 'income', 'expense', 'budget',
            'investment', 'portfolio', 'asset', 'liability', 'equity', 'balance sheet',
            'income statement', 'cash flow', 'audit', 'tax'
        ]
        
        technical_keywords = [
            'technical', 'specification', 'requirement', 'system', 'software', 'hardware',
            'architecture', 'design', 'implementation', 'api', 'database', 'protocol',
            'algorithm', 'framework', 'development', 'testing'
        ]
        
        medical_keywords = [
            'medical', 'health', 'patient', 'diagnosis', 'treatment', 'symptom',
            'medication', 'prescription', 'doctor', 'physician', 'clinic', 'hospital',
            'therapy', 'recovery', 'prognosis', 'medical record'
        ]
        
        # Check filename patterns
        if any(keyword in filename_lower for keyword in ['policy', 'insurance', 'coverage']):
            return "Policy Wordings"
        elif any(keyword in filename_lower for keyword in ['contract', 'agreement', 'legal']):
            return "Legal Documents"
        elif any(keyword in filename_lower for keyword in ['financial', 'report', 'statement']):
            return "Financial Documents"
        elif any(keyword in filename_lower for keyword in ['technical', 'spec', 'manual']):
            return "Technical Documents"
        elif any(keyword in filename_lower for keyword in ['medical', 'health', 'patient']):
            return "Medical Documents"
        
        # Check content patterns
        insurance_score = sum(1 for keyword in insurance_keywords if keyword in text_lower)
        legal_score = sum(1 for keyword in legal_keywords if keyword in text_lower)
        financial_score = sum(1 for keyword in financial_keywords if keyword in text_lower)
        technical_score = sum(1 for keyword in technical_keywords if keyword in text_lower)
        medical_score = sum(1 for keyword in medical_keywords if keyword in text_lower)
        
        # Determine document type based on highest score
        scores = {
            "Policy Wordings": insurance_score,
            "Legal Documents": legal_score,
            "Financial Documents": financial_score,
            "Technical Documents": technical_score,
            "Medical Documents": medical_score
        }
        
        max_score = max(scores.values())
        if max_score >= 3:  # Threshold for confident detection
            return max(scores, key=scores.get)
        else:
            return "unknown"
    
    def process_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process PDF file and extract text"""
        try:
            # Convert bytes to BytesIO object for pypdf
            pdf_file = io.BytesIO(file_content)
            pdf_reader = pypdf.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return {
                "text": text,
                "pages": len(pdf_reader.pages),
                "filename": filename
            }
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
    
    def process_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process DOCX file and extract text"""
        try:
            # Convert bytes to BytesIO object for python-docx
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return {
                "text": text,
                "paragraphs": len(doc.paragraphs),
                "filename": filename
            }
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")
    
    def process_email(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process email file and extract text"""
        try:
            msg = email.message_from_bytes(file_content, policy=policy.default)
            text = ""
            
            # Extract subject
            subject = msg.get('subject', '')
            text += f"Subject: {subject}\n\n"
            
            # Extract body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        text += part.get_payload(decode=True).decode()
                        break
            else:
                text += msg.get_payload(decode=True).decode()
            
            return {
                "text": text,
                "subject": subject,
                "filename": filename
            }
        except Exception as e:
            raise Exception(f"Error processing email: {str(e)}")
    
    def chunk_text(self, text: str, document_id: str) -> List[LangchainDocument]:
        """Split text into chunks using LangChain"""
        try:
            chunks = self.text_splitter.split_text(text)
            documents = []
            
            for i, chunk in enumerate(chunks):
                doc = LangchainDocument(
                    page_content=chunk,
                    metadata={
                        "document_id": document_id,
                        "chunk_id": i,
                        "chunk_index": i,
                        "total_chunks": len(chunks)
                    }
                )
                documents.append(doc)
            
            return documents
        except Exception as e:
            raise Exception(f"Error chunking text: {str(e)}")
    
    def process_document(self, file_content: bytes, filename: str, file_type: str) -> Dict[str, Any]:
        """Main method to process any document type"""
        document_id = str(uuid.uuid4())
        
        # Process based on file type
        if file_type.lower() == "pdf":
            result = self.process_pdf(file_content, filename)
        elif file_type.lower() in ["docx", "doc"]:
            result = self.process_docx(file_content, filename)
        elif file_type.lower() in ["eml", "email"]:
            result = self.process_email(file_content, filename)
        else:
            raise Exception(f"Unsupported file type: {file_type}")
        
        # Detect document type automatically
        detected_type = self.detect_document_type(result["text"], filename)
        
        # Chunk the text
        chunks = self.chunk_text(result["text"], document_id)
        
        return {
            "document_id": document_id,
            "filename": filename,
            "file_type": file_type,
            "document_type": detected_type,
            "text": result["text"],
            "chunks": chunks,
            "total_chunks": len(chunks),
            "upload_time": datetime.now()
        } 